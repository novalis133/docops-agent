"""Document parser for PDF, DOCX, and Markdown files."""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
import markdown
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError


@dataclass
class ParsedSection:
    """Represents a parsed section of a document."""

    title: str
    content: str
    level: int = 1
    page_number: Optional[int] = None
    metadata: dict = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """Represents a fully parsed document."""

    filename: str
    file_type: str
    title: str
    sections: list[ParsedSection]
    raw_text: str
    metadata: dict = field(default_factory=dict)
    page_count: int = 1


class DocumentParser:
    """Parse documents from various formats into structured sections."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".md", ".markdown", ".txt"}

    def __init__(self) -> None:
        self._heading_pattern = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)

    def parse(self, file_path: str | Path) -> ParsedDocument:
        """Parse a document from the given file path."""
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}")

        if extension == ".pdf":
            return self._parse_pdf(path)
        elif extension in {".docx", ".doc"}:
            return self._parse_docx(path)
        elif extension in {".md", ".markdown"}:
            return self._parse_markdown(path)
        elif extension == ".txt":
            return self._parse_text(path)

        raise ValueError(f"Unsupported file type: {extension}")

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        """Parse a PDF document using PyMuPDF."""
        sections: list[ParsedSection] = []
        raw_text_parts: list[str] = []
        title = path.stem

        try:
            doc = fitz.open(path)
            page_count = len(doc)

            for page_num, page in enumerate(doc, start=1):
                text = page.get_text()
                raw_text_parts.append(text)

                # Extract sections based on font size analysis
                blocks = page.get_text("dict")["blocks"]
                current_section_title = f"Page {page_num}"
                current_section_content: list[str] = []

                for block in blocks:
                    if "lines" not in block:
                        continue

                    for line in block["lines"]:
                        line_text = ""
                        max_font_size = 0

                        for span in line["spans"]:
                            line_text += span["text"]
                            max_font_size = max(max_font_size, span["size"])

                        line_text = line_text.strip()
                        if not line_text:
                            continue

                        # Heuristic: larger fonts are likely headings
                        if max_font_size > 14 and len(line_text) < 100:
                            # Save previous section
                            if current_section_content:
                                sections.append(
                                    ParsedSection(
                                        title=current_section_title,
                                        content="\n".join(current_section_content),
                                        level=1,
                                        page_number=page_num,
                                    )
                                )
                            current_section_title = line_text
                            current_section_content = []
                        else:
                            current_section_content.append(line_text)

                # Add remaining content as a section
                if current_section_content:
                    sections.append(
                        ParsedSection(
                            title=current_section_title,
                            content="\n".join(current_section_content),
                            level=1,
                            page_number=page_num,
                        )
                    )

            doc.close()

            # Extract title from first heading if available
            if sections and sections[0].title != "Page 1":
                title = sections[0].title

            return ParsedDocument(
                filename=path.name,
                file_type="pdf",
                title=title,
                sections=sections if sections else [
                    ParsedSection(
                        title="Content",
                        content="\n".join(raw_text_parts),
                        level=1,
                    )
                ],
                raw_text="\n\n".join(raw_text_parts),
                page_count=page_count,
                metadata={"source": str(path)},
            )

        except Exception as e:
            raise RuntimeError(f"Failed to parse PDF: {e}") from e

    def _parse_docx(self, path: Path) -> ParsedDocument:
        """Parse a DOCX document."""
        sections: list[ParsedSection] = []
        raw_text_parts: list[str] = []
        title = path.stem

        try:
            doc = DocxDocument(path)

            current_section_title = "Introduction"
            current_section_content: list[str] = []
            current_level = 1

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                raw_text_parts.append(text)
                style_name = para.style.name if para.style else ""

                # Check if this is a heading
                if style_name.startswith("Heading"):
                    # Save previous section
                    if current_section_content:
                        sections.append(
                            ParsedSection(
                                title=current_section_title,
                                content="\n".join(current_section_content),
                                level=current_level,
                            )
                        )

                    # Extract heading level
                    level_match = re.search(r"Heading\s*(\d+)", style_name)
                    current_level = int(level_match.group(1)) if level_match else 1
                    current_section_title = text
                    current_section_content = []

                    # Use first heading as document title
                    if not sections and current_level == 1:
                        title = text
                elif style_name == "Title":
                    title = text
                else:
                    current_section_content.append(text)

            # Add final section
            if current_section_content:
                sections.append(
                    ParsedSection(
                        title=current_section_title,
                        content="\n".join(current_section_content),
                        level=current_level,
                    )
                )

            return ParsedDocument(
                filename=path.name,
                file_type="docx",
                title=title,
                sections=sections if sections else [
                    ParsedSection(
                        title="Content",
                        content="\n".join(raw_text_parts),
                        level=1,
                    )
                ],
                raw_text="\n\n".join(raw_text_parts),
                metadata={"source": str(path)},
            )

        except PackageNotFoundError as e:
            raise FileNotFoundError(f"Invalid DOCX file: {e}") from e
        except Exception as e:
            raise RuntimeError(f"Failed to parse DOCX: {e}") from e

    def _parse_markdown(self, path: Path) -> ParsedDocument:
        """Parse a Markdown document."""
        content = path.read_text(encoding="utf-8")
        sections: list[ParsedSection] = []
        title = path.stem

        # Convert to HTML for easier parsing
        html = markdown.markdown(content, extensions=["tables", "fenced_code"])
        soup = BeautifulSoup(html, "html.parser")

        # Extract plain text
        raw_text = soup.get_text()

        # Parse sections from headings
        current_section_title = "Introduction"
        current_section_content: list[str] = []
        current_level = 1

        for element in soup.children:
            if element.name and element.name.startswith("h") and len(element.name) == 2:
                # Save previous section
                if current_section_content:
                    sections.append(
                        ParsedSection(
                            title=current_section_title,
                            content="\n".join(current_section_content),
                            level=current_level,
                        )
                    )

                current_level = int(element.name[1])
                current_section_title = element.get_text().strip()
                current_section_content = []

                # First h1 is the title
                if not sections and current_level == 1:
                    title = current_section_title
            elif element.name:
                text = element.get_text().strip()
                if text:
                    current_section_content.append(text)

        # Add final section
        if current_section_content:
            sections.append(
                ParsedSection(
                    title=current_section_title,
                    content="\n".join(current_section_content),
                    level=current_level,
                )
            )

        return ParsedDocument(
            filename=path.name,
            file_type="markdown",
            title=title,
            sections=sections if sections else [
                ParsedSection(title="Content", content=raw_text, level=1)
            ],
            raw_text=raw_text,
            metadata={"source": str(path)},
        )

    def _parse_text(self, path: Path) -> ParsedDocument:
        """Parse a plain text document."""
        content = path.read_text(encoding="utf-8")

        # Try to detect sections by blank lines or patterns
        sections: list[ParsedSection] = []
        paragraphs = re.split(r"\n\s*\n", content)

        for i, para in enumerate(paragraphs):
            para = para.strip()
            if para:
                sections.append(
                    ParsedSection(
                        title=f"Section {i + 1}",
                        content=para,
                        level=1,
                    )
                )

        return ParsedDocument(
            filename=path.name,
            file_type="text",
            title=path.stem,
            sections=sections if sections else [
                ParsedSection(title="Content", content=content, level=1)
            ],
            raw_text=content,
            metadata={"source": str(path)},
        )

    def parse_bytes(
        self, content: bytes, filename: str, file_type: str
    ) -> ParsedDocument:
        """Parse document from bytes content."""
        import tempfile

        # Write to temporary file and parse
        suffix = f".{file_type}" if not file_type.startswith(".") else file_type

        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(content)
            tmp_path = Path(tmp.name)

        try:
            doc = self.parse(tmp_path)
            doc.filename = filename
            return doc
        finally:
            tmp_path.unlink(missing_ok=True)
